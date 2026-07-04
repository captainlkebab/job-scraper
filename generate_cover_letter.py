import os
import json
import re
from datetime import datetime
from dotenv import load_dotenv
from groq import Groq
from docx import Document
from docx.shared import Pt, Inches
from docx2pdf import convert

load_dotenv()

def analyze_job_with_groq(job_title, job_description):
    """
    Nutzt Groq, um basierend auf dem Bewerberprofil und der Anzeige
    ein perfekt abgestimmtes, wiederholungsfreies Anschreiben zu generieren.
    """
    print(f"🤖 Groq generiert das maßgeschneiderte Anschreiben für: '{job_title}'...")
    
    client = Groq()
    
    user_profile = """
    Bewerber-Profil (Samil Demiroglu):
    - Aktueller Status: Masterstudent Business Data Science an der Aalborg Universität (AAU).
    - Technische Kernkompetenzen: Python, SQL, TypeScript, React, React Native.
    - Fundierte Projekterfahrung: 
      * Entwicklung und Bereitstellung von automatisierten, wöchentlichen Web-Scrapern zur Datenaggregation (Preisscraper für >2.500 Produkte).
      * Konzeption komplexer, relationaler Datenbankmodelle und fortgeschrittenes Data Engineering.
      * Praktischer Einsatz von Large Language Models (LLMs) zur Strukturierung unstrukturierter Textdaten und Aufdeckung verborgener Datenbeziehungen.
    - Fokus für die Masterarbeit: Anwendung von Data Science, KI und Prozessoptimierung/Automatisierung im unternehmerischen Kontext.
    - Sprachkenntnisse: Deutsch (Muttersprache), Englisch (C2 / verhandlungssicher).
    """

    prompt = f"""
    Du bist ein präziser HR-Experte. Erstelle den Text für ein professionelles Anschreiben für eine Masterarbeit.
    Nutze Samils Profil und stimme den Text dynamisch auf die Anforderungen der Stellenanzeige ab.

    {user_profile}

    Stellentitel: {job_title}
    Stellenbeschreibung:
    {job_description}

    --- ARCHITEKTUR DES ANSCHREIBENS ---
    Das Anschreiben wird im Skript wie folgt aufgebaut. Du musst die passenden Textbausteine liefern.
    Achte penibel darauf, dass sich Informationen (wie Samils Status als AAU-Student) NICHT im Hauptteil wiederholen, da dies bereits in der Einleitung steht!

    [Einleitung]: "mit großem Interesse habe ich Ihre Ausschreibung für eine Masterarbeit im Bereich Data Science gelesen. Als Masterstudent der Fachrichtung Business Data Science bringe ich fundierte Kenntnisse in Python, SQL und der automatisierten Datenverarbeitung mit. Besonders reizt mich an der ausgeschriebenen Stelle bei [unternehmensname] die Möglichkeit, an [konkrete_aufgabe]."
    [Hauptteil / profil_match_absatz]: (Hier schlägst du die Brücke zu Samils Scraper-, Datenbank- oder LLM-Erfahrung, passend zur Stelle).
    [Schluss]: "Bezüglich des konkreten Projektschwerpunkts bin ich fachlich offen, da mich die methodische Bandbreite moderner Data-Science-Infrastrukturen fasziniert. Ich weiß, dass das Deployment und die Skalierung von Modellen unter industriellen Bedingungen neue architektonische Herausforderungen mit sich bringen. Genau auf diese steile Lernkurve und die Zusammenarbeit im Team freue ich mich. Über die Gelegenheit, meine Motivation in einem persönlichen Gespräch vorzustellen, freue ich mich sehr."

    Gib AUSSCHLIESSLICH ein valides JSON-Objekt mit folgenden Keys zurück:
    1. "unternehmensname": Name der Firma (z.B. "Mercedes-Benz AG" oder "HELLA"). Ohne Rechtsformzusätze wie GmbH, wenn es unnatürlich wirkt.
    2. "bereinigter_titel": Der originale Stellentitel, aber absolut bereinigt um Zusätze wie (m/w/d), (w/m/div), Ref-Nummern oder Standortangaben. (Beispiel: "Masterand: Automatisierte Sensorkalibrierung und Fehlerklassifizierung").
    3. "konkrete_aufgabe": Ein präziser Nebensatz, der perfekt hinter "die Möglichkeit, an..." passt (z.B. "der Entwicklung datengestützter Automatisierungslösungen mitzuwirken").
    4. "ansprechpartner_zeile": Der vollständige Name der Ansprechperson für die Adresse (z.B. "Frau Dr. Martina Müller"). Wenn absolut kein Name in der Anzeige zu finden ist, gib einen LEEREN STRING "" zurück.
    5. "anrede_formel": "Sehr geehrte/r Frau/Herr [Nachname]," oder "Sehr geehrte Damen und Herren," falls kein Ansprechpartner existiert.
    6. "profil_match_absatz": Ein starker, maßgeschneiderter Absatz (3-4 Sätze) für den Hauptteil. Verbinde Samils Tech-Stack/Erfahrung mit der Stelle. KEINE Floskeln, keine em-dashes (—), KEINE Wiederholung der Universität oder des Studiengangs!
    """

    chat_completion = client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model="gpt-oss-120b",
        response_format={"type": "json_object"}
    )

    return json.loads(chat_completion.choices[0].message.content)


def create_cover_letter_files(job_title, job_description, base_output_dir="coverLetters"):
    date_suffix = datetime.now().strftime("%y-%m-%d")
    output_dir = os.path.join(base_output_dir, f"CL_{date_suffix}")
    
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    try:
        llm_data = analyze_job_with_groq(job_title, job_description)
    except Exception as e:
        print(f"❌ Groq-API-Fehler: {e}")
        return None

    safe_company = re.sub(r'[\\/*?:"<>| ]', "_", llm_data.get("unternehmensname", "Firma"))
    safe_title = re.sub(r'[\\/*?:"<>| ]', "_", llm_data.get("bereinigter_titel", job_title))[:20]
    
    docx_path = f"{output_dir}/Anschreiben_{safe_company}_{safe_title}.docx"
    pdf_path = f"{output_dir}/Anschreiben_{safe_company}_{safe_title}.pdf"

    doc = Document()
    
    # Seitenränder setzen
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.15

    # Absender
    p_sender = doc.add_paragraph()
    p_sender.add_run("Samil Demiroglu\nKorfesstraße 37\n38104 Braunschweig").font.size = Pt(10)
    p_sender.paragraph_format.space_after = Pt(36)

    # Empfänger (Dynamische Zeilenanzahl je nach Ansprechpartner)
    p_company = doc.add_paragraph()
    company_text = f"{llm_data.get('unternehmensname', 'Ihrem Unternehmen')}\n"
    if llm_data.get("ansprechpartner_zeile"):
        company_text += f"{llm_data.get('ansprechpartner_zeile')}\n"
    company_text += "Personalabteilung"
    p_company.add_run(company_text)
    p_company.paragraph_format.space_after = Pt(24)

    # Datum
    p_date = doc.add_paragraph()
    p_date.alignment = 2
    p_date.add_run(f"Braunschweig, den {datetime.now().strftime('%d.%m.%Y')}")
    p_date.paragraph_format.space_after = Pt(24)

    # --- BETREFFZEILE (Clean & professionell ohne \n) ---
    p_subject = doc.add_paragraph()
    subject_text = f"Bewerbung auf die Stelle {llm_data.get('bereinigter_titel', job_title)}"
    subject_run = p_subject.add_run(subject_text)
    subject_run.bold = True
    subject_run.font.size = Pt(11)
    p_subject.paragraph_format.space_after = Pt(20)

    # Anrede
    p_salutation = doc.add_paragraph(llm_data.get('anrede_formel', 'Sehr geehrte Damen und Herren,'))
    p_salutation.paragraph_format.space_after = Pt(12)

    paragraph_space = Pt(10)

    # Einleitung
    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_after = paragraph_space
    p_intro.add_run("mit großem Interesse habe ich Ihre Ausschreibung für eine Masterarbeit im Bereich Data Science gelesen. Als Masterstudent der Fachrichtung Business Data Science bringe ich fundierte Kenntnisse in Python, SQL und der automatisierten Datenverarbeitung mit. Besonders reizt mich an der ausgeschriebenen Stelle bei ")
    p_intro.add_run(llm_data.get('unternehmensname', 'Ihrem Unternehmen'))
    p_intro.add_run(" die Möglichkeit, an ")
    p_intro.add_run(llm_data.get('konkrete_aufgabe', 'den anstehenden Projekten zu arbeiten'))
    p_intro.add_run(".")

    # KI-Match-Absatz (Garantiert ohne Redundanzen)
    p_match = doc.add_paragraph(llm_data.get('profil_match_absatz', ''))
    p_match.paragraph_format.space_after = paragraph_space

    # Neuer, knackigerer Schlussteil
    p_close = doc.add_paragraph()
    p_close.paragraph_format.space_after = paragraph_space
    p_close.add_run( 
        "Bezüglich des konkreten Projektschwerpunkts bin ich fachlich offen, da mich die methodische Bandbreite moderner Data-Science-Infrastrukturen fasziniert. Ich weiß, dass das Deployment und die Skalierung von Modellen unter industriellen Bedingungen neue architektonische Herausforderungen mit sich bringen. Genau auf diese steile Lernkurve und die Zusammenarbeit im Team freue ich mich."
    )
    
    p_call_to_action = doc.add_paragraph("Über die Gelegenheit, meine Motivation in einem persönlichen Gespräch vorzustellen, freue ich mich sehr.")
    p_call_to_action.paragraph_format.space_after = Pt(24)

    doc.add_paragraph("Mit freundlichen Grüßen\n\n\nSamil Demiroglu")

    doc.save(docx_path)
    print(f"💾 DOCX gespeichert: {docx_path}")

    try:
        print("🖨️ Generiere PDF...")
        convert(docx_path, pdf_path)
        print(f"✅ PDF erfolgreich exportiert: {pdf_path}")
    except Exception as e:
        print(f"⚠️ PDF-Export fehlgeschlagen: {e}")

    return pdf_path


# --- AUTOMATISIERTER LOOP ÜBER RELEVANTE STELLEN ---
if __name__ == "__main__":
    date_str = datetime.now().strftime("%y-%m-%d")
    test_json_path = f"jobs_output/indeed_{date_str}.json"
    
    if os.path.exists(test_json_path):
        with open(test_json_path, "r", encoding="utf-8") as f:
            jobs = json.load(f)
        
        if jobs:
            # Regex für die relevanten Keywords (Masterarbeit, Masterand, Master Thesis, Abschlussarbeit)
            title_filter = re.compile(r'(masterarbeit|masterand|master\s*thesis|abschlussarbeit)', re.IGNORECASE)
            match_count = 0
            
            print(f"🚀 Starte Filterung über {len(jobs)} extrahierte Stellen...")
            
            for job in jobs:
                job_title = job.get('title', 'Unbekannter Titel')
                
                # Wenn die Keywords matchen, generieren wir das Anschreiben
                if title_filter.search(job_title):
                    match_count += 1
                    print(f"\n🎯 [Treffer #{match_count}] Verarbeite: '{job_title}'")
                    
                    try:
                        create_cover_letter_files(job_title, job.get('description', ''))
                    except Exception as e:
                        print(f"❌ Fehler bei dieser Stelle: {e}")
                        continue
            
            print(f"\n✅ Fertig! Insgesamt wurden {match_count} passende Anschreiben generiert.")
        else:
            print("ℹ️ Die JSON-Datei ist leer.")
    else:
        print(f"ℹ️ Test-JSON unter '{test_json_path}' nicht gefunden. Bitte Scraper vorher ausführen.")